from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "artifacts" / "stage1_documents"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOC_SKILL = Path(
    r"C:/Users/bdq/.codex/plugins/cache/openai-primary-runtime/documents/26.715.12143/skills/documents"
)
sys.path.insert(0, str(DOC_SKILL / "scripts"))
from table_geometry import apply_table_geometry  # noqa: E402


PLAYER_GUIDE = OUT_DIR / "零网寻踪_玩家指南与玩法任务_阶段1.docx"
WORLD_BOOK = OUT_DIR / "零网寻踪_世界观设定手册_阶段1.docx"

FONT_BODY = "Microsoft YaHei"
FONT_MONO = "Consolas"

INK = "20211E"
MUTED = "686B63"
PAPER = "F1ECD9"
PAPER_LIGHT = "F8F5E9"
ARCHIVE = "505840"
ARCHIVE_DARK = "252B25"
CRT_GREEN = "33D968"
CRT_GREEN_BRIGHT = "33FF66"
AMBER = "D88A00"
RED = "A7352B"
RED_DARK = "6E201B"
BLUE_GRAY = "DDE4E1"
GRID = "8D9188"
WHITE = "FFFFFF"
BLACK = "0A0C0B"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name=FONT_BODY, size=None, color=INK, bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr()
    r_fonts = run._element.rPr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), name)
    r_fonts.set(qn("w:hAnsi"), name)
    r_fonts.set(qn("w:eastAsia"), FONT_BODY if name == FONT_MONO else name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def set_cell_border(cell, **edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge_name, edge_data in edges.items():
        tag = "w:" + edge_name
        edge = tc_borders.find(qn(tag))
        if edge is None:
            edge = OxmlElement(tag)
            tc_borders.append(edge)
        for key, value in edge_data.items():
            edge.set(qn("w:" + key), str(value))


def shade_paragraph(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")


def paragraph_bottom_border(paragraph, color=GRID, size=8, space=3):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = p_bdr.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        p_bdr.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_text(cell, text, *, size=9.6, color=INK, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT, mono=False):
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.clear()
    run = p.add_run(str(text))
    set_run_font(run, FONT_MONO if mono else FONT_BODY, size=size, color=color, bold=bold)
    return p


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = instruction
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr_text)
    run._r.append(fld_char2)
    return run


def set_document_basics(doc: Document, doc_code: str, running_title: str):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.86)
    section.right_margin = Inches(0.86)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_BODY
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in [
        ("Title", 27, ARCHIVE_DARK, 0, 8),
        ("Subtitle", 12.5, MUTED, 0, 12),
        ("Heading 1", 17, ARCHIVE_DARK, 15, 8),
        ("Heading 2", 13.5, ARCHIVE, 11, 5),
        ("Heading 3", 11.5, RED_DARK, 8, 3),
    ]:
        st = styles[name]
        st.font.name = FONT_BODY
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        st.font.size = Pt(size)
        st.font.color.rgb = rgb(color)
        st.font.bold = name != "Subtitle"
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    for name in ["List Bullet", "List Number"]:
        st = styles[name]
        st.font.name = FONT_BODY
        st._element.rPr.rFonts.set(qn("w:ascii"), FONT_BODY)
        st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_BODY)
        st._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        st.font.size = Pt(10.5)
        st.paragraph_format.left_indent = Inches(0.38)
        st.paragraph_format.first_line_indent = Inches(-0.19)
        st.paragraph_format.space_after = Pt(4)
        st.paragraph_format.line_spacing = 1.18

    if "Archive Label" not in [s.name for s in styles]:
        label = styles.add_style("Archive Label", WD_STYLE_TYPE.PARAGRAPH)
        label.font.name = FONT_MONO
        label._element.rPr.rFonts.set(qn("w:ascii"), FONT_MONO)
        label._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_MONO)
        label._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_BODY)
        label.font.size = Pt(8.6)
        label.font.bold = True
        label.font.color.rgb = rgb(MUTED)
        label.paragraph_format.space_before = Pt(0)
        label.paragraph_format.space_after = Pt(3)
        label.paragraph_format.keep_with_next = True

    doc.core_properties.title = running_title
    doc.core_properties.subject = "《零网寻踪》阶段 1 文档"
    doc.core_properties.author = "零网寻踪项目组"
    doc.core_properties.keywords = "零网寻踪, 玩家指南, 世界观, 主持说明, 阶段1"

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(2)
    r1 = hp.add_run(f"{running_title}  //  ")
    set_run_font(r1, FONT_MONO, size=8.2, color=MUTED, bold=True)
    r2 = hp.add_run(doc_code)
    set_run_font(r2, FONT_MONO, size=8.2, color=RED, bold=True)
    paragraph_bottom_border(hp, color=GRID, size=4, space=2)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = fp.add_run("ARCHIVE COPY  |  ")
    set_run_font(r, FONT_MONO, size=8, color=MUTED)
    add_field(fp, "PAGE")
    r2 = fp.add_run("  |  VOID IF DETACHED")
    set_run_font(r2, FONT_MONO, size=8, color=MUTED)


def add_cover(doc: Document, *, code: str, title: str, subtitle: str, deck: str, edition: str, audience: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f"FILE // {code}")
    set_run_font(r, FONT_MONO, size=9, color=RED, bold=True)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    apply_table_geometry(table, [9380], indent_dxa=360, cell_margins_dxa={"top": 320, "bottom": 320, "start": 360, "end": 360})
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, top={"val": "single", "sz": 18, "color": CRT_GREEN}, bottom={"val": "single", "sz": 18, "color": CRT_GREEN}, start={"val": "single", "sz": 18, "color": CRT_GREEN}, end={"val": "single", "sz": 18, "color": CRT_GREEN})

    p0 = cell.paragraphs[0]
    p0.paragraph_format.space_after = Pt(50)
    r0 = p0.add_run("ZERO-NET ARCHIVE / SIGNAL RECOVERED")
    set_run_font(r0, FONT_MONO, size=9.5, color=CRT_GREEN, bold=True)

    pt = cell.add_paragraph()
    pt.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pt.paragraph_format.space_after = Pt(8)
    rt = pt.add_run(title)
    set_run_font(rt, FONT_BODY, size=29, color=CRT_GREEN_BRIGHT, bold=True)

    ps = cell.add_paragraph()
    ps.paragraph_format.space_after = Pt(22)
    rs = ps.add_run(subtitle)
    set_run_font(rs, FONT_BODY, size=15, color=AMBER, bold=True)

    pd = cell.add_paragraph()
    pd.paragraph_format.space_after = Pt(52)
    rd = pd.add_run(deck)
    set_run_font(rd, FONT_BODY, size=11.2, color=WHITE)

    stamp = cell.add_paragraph()
    stamp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    stamp.paragraph_format.space_after = Pt(34)
    rstamp = stamp.add_run("  机密 / 仅限档案持有人  ")
    set_run_font(rstamp, FONT_BODY, size=10, color=WHITE, bold=True)
    shade_paragraph(stamp, RED_DARK)

    meta = cell.add_paragraph()
    meta.paragraph_format.space_after = Pt(6)
    rm = meta.add_run(f"EDITION : {edition}\nAUDIENCE: {audience}\nSTATUS  : VERIFIED STAGE-1 BUILD\nSIGNAL  : REN-0 / TRACE PENDING")
    set_run_font(rm, FONT_MONO, size=9.2, color="B7C2B9")

    last = cell.add_paragraph()
    last.paragraph_format.space_before = Pt(30)
    last.paragraph_format.space_after = Pt(0)
    rl = last.add_run("DO NOT TRUST THE SUMMARY. KEEP THE ORIGINAL TIMESTAMP.")
    set_run_font(rl, FONT_MONO, size=8.6, color=CRT_GREEN)

    doc.add_page_break()


def add_archive_kicker(doc, text, color=RED):
    p = doc.add_paragraph(style="Archive Label")
    r = p.add_run(text.upper())
    set_run_font(r, FONT_MONO, size=8.5, color=color, bold=True)
    return p


def add_section_title(doc, number: str, title: str, deck: str | None = None, *, page_break=False):
    kicker = add_archive_kicker(doc, f"SECTION {number}")
    kicker.paragraph_format.page_break_before = page_break
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(title)
    set_run_font(r, FONT_BODY, size=17, color=ARCHIVE_DARK, bold=True)
    paragraph_bottom_border(p, color=ARCHIVE, size=9, space=4)
    if deck:
        d = doc.add_paragraph()
        d.paragraph_format.space_after = Pt(9)
        rr = d.add_run(deck)
        set_run_font(rr, FONT_BODY, size=10.5, color=MUTED, italic=True)


def add_h2(doc, title: str):
    p = doc.add_paragraph(style="Heading 2")
    r = p.add_run(title)
    set_run_font(r, FONT_BODY, size=13.5, color=ARCHIVE, bold=True)
    return p


def add_h3(doc, title: str):
    p = doc.add_paragraph(style="Heading 3")
    r = p.add_run(title)
    set_run_font(r, FONT_BODY, size=11.5, color=RED_DARK, bold=True)
    return p


def add_body(doc, text: str, *, bold_prefix: str | None = None, color=INK, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, FONT_BODY, size=10.5, color=color, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2, FONT_BODY, size=10.5, color=color, italic=italic)
    else:
        r = p.add_run(text)
        set_run_font(r, FONT_BODY, size=10.5, color=color, italic=italic)
    return p


def _new_single_level_num_id(doc, *, num_format: str, level_text: str):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(el.get(qn("w:abstractNumId")))
        for el in numbering.findall(qn("w:abstractNum"))
        if el.get(qn("w:abstractNumId")) is not None
    ]
    num_ids = [
        int(el.get(qn("w:numId")))
        for el in numbering.findall(qn("w:num"))
        if el.get(qn("w:numId")) is not None
    ]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    nsid = OxmlElement("w:nsid")
    nsid.set(qn("w:val"), f"A5{abstract_id:06X}"[-8:])
    abstract.append(nsid)
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    tmpl = OxmlElement("w:tmpl")
    tmpl.set(qn("w:val"), f"B7{abstract_id:06X}"[-8:])
    abstract.append(tmpl)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), num_format)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), level_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    p_pr = OxmlElement("w:pPr")
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "540")
    ind.set(qn("w:hanging"), "270")
    p_pr.append(ind)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "80")
    spacing.set(qn("w:line"), "300")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    lvl.extend([start, num_fmt, lvl_text, lvl_jc, p_pr])
    abstract.append(lvl)
    first_num = numbering.find(qn("w:num"))
    if first_num is None:
        numbering.append(abstract)
    else:
        numbering.insert(numbering.index(first_num), abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    lvl_override = OxmlElement("w:lvlOverride")
    lvl_override.set(qn("w:ilvl"), "0")
    start_override = OxmlElement("w:startOverride")
    start_override.set(qn("w:val"), "1")
    lvl_override.append(start_override)
    num.append(lvl_override)
    numbering.append(num)
    return num_id


def _new_decimal_num_id(doc):
    return _new_single_level_num_id(doc, num_format="decimal", level_text="%1.")


def _new_bullet_num_id(doc):
    return _new_single_level_num_id(doc, num_format="bullet", level_text="•")


def _apply_num_id(paragraph, num_id):
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num_id_el])
    paragraph._p.get_or_add_pPr().append(num_pr)


def add_bullets(doc, items, level=0, *, compact=False):
    num_id = _new_bullet_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        _apply_num_id(p, num_id)
        if level:
            p.paragraph_format.left_indent = Inches(0.38 + 0.28 * level)
        p.paragraph_format.space_after = Pt(1.5 if compact else 4)
        p.paragraph_format.line_spacing = 1.12 if compact else 1.25
        r = p.add_run(item)
        set_run_font(r, FONT_BODY, size=9.8 if compact else 10.2, color=INK)


def add_numbered(doc, items):
    num_id = _new_decimal_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        _apply_num_id(p, num_id)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(item)
        set_run_font(r, FONT_BODY, size=10.2, color=INK)


def add_callout(doc, label: str, text: str, *, fill=BLUE_GRAY, accent=ARCHIVE, text_color=INK):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9560], indent_dxa=180, cell_margins_dxa={"top": 130, "bottom": 130, "start": 180, "end": 180})
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, start={"val": "single", "sz": 24, "color": accent}, top={"val": "nil"}, bottom={"val": "nil"}, end={"val": "nil"})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    rl = p.add_run(label + "  ")
    set_run_font(rl, FONT_MONO, size=8.8, color=accent, bold=True)
    rt = p.add_run(text)
    set_run_font(rt, FONT_BODY, size=10, color=text_color)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_code_block(doc, lines, label="TERMINAL EXTRACT"):
    table = doc.add_table(rows=1, cols=1)
    apply_table_geometry(table, [9550], indent_dxa=190, cell_margins_dxa={"top": 150, "bottom": 150, "start": 190, "end": 190})
    cell = table.cell(0, 0)
    set_cell_shading(cell, BLACK)
    set_cell_border(cell, top={"val": "single", "sz": 8, "color": CRT_GREEN}, bottom={"val": "single", "sz": 8, "color": CRT_GREEN}, start={"val": "single", "sz": 8, "color": CRT_GREEN}, end={"val": "single", "sz": 8, "color": CRT_GREEN})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(label)
    set_run_font(r, FONT_MONO, size=8, color=AMBER, bold=True)
    for line in lines:
        pl = cell.add_paragraph()
        pl.paragraph_format.space_after = Pt(0)
        rr = pl.add_run(line)
        set_run_font(rr, FONT_MONO, size=8.8, color=CRT_GREEN)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_table(doc, headers, rows, widths, *, header_fill=ARCHIVE_DARK, font_size=9.2, first_col_bold=False, alignments=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, size=8.9, color=WHITE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], header_fill)
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for c_idx, value in enumerate(row):
            alignment = alignments[c_idx] if alignments else WD_ALIGN_PARAGRAPH.LEFT
            set_cell_text(cells[c_idx], value, size=font_size, color=INK, bold=(first_col_bold and c_idx == 0), align=alignment, mono=False)
            if r_idx % 2 == 1:
                set_cell_shading(cells[c_idx], PAPER_LIGHT)
    source_total = sum(widths)
    fitted = [int(round(width * 9620 / source_total)) for width in widths]
    fitted[-1] += 9620 - sum(fitted)
    apply_table_geometry(table, fitted, indent_dxa=120, cell_margins_dxa={"top": 95, "bottom": 95, "start": 120, "end": 120})
    for row in table.rows:
        row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    return table


def add_checklist(doc, items, *, widths=(520, 9220)):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for idx, item in enumerate(items):
        cells = table.add_row().cells
        set_cell_text(cells[0], "□", size=13, color=ARCHIVE, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], item, size=9.8, color=INK)
        if idx % 2 == 1:
            set_cell_shading(cells[0], PAPER_LIGHT)
            set_cell_shading(cells[1], PAPER_LIGHT)
    source_total = sum(widths)
    fitted = [int(round(width * 9620 / source_total)) for width in widths]
    fitted[-1] += 9620 - sum(fitted)
    apply_table_geometry(table, fitted, indent_dxa=120, cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_image(doc, image_path: Path, caption: str, width=6.68):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width))
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.paragraph_format.space_after = Pt(8)
    r = cp.add_run(caption)
    set_run_font(r, FONT_BODY, size=8.5, color=MUTED, italic=True)


def add_mission_header(doc, tag: str, title: str, time_box: str, question: str, *, page_break=False):
    kicker = add_archive_kicker(doc, tag, color=RED)
    kicker.paragraph_format.page_break_before = page_break
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, FONT_BODY, size=23, color=ARCHIVE_DARK, bold=True)
    info = doc.add_paragraph()
    info.paragraph_format.space_after = Pt(8)
    ri = info.add_run(time_box)
    set_run_font(ri, FONT_MONO, size=9, color=AMBER, bold=True)
    add_callout(doc, "核心问题", question, fill="E6E1D1", accent=RED_DARK)


def add_spoiler_band(doc, label="主持/解谜专用 · 以下含完整答案", *, page_break=False):
    p = doc.add_paragraph()
    p.paragraph_format.page_break_before = page_break
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    shade_paragraph(p, RED_DARK)
    r = p.add_run(label)
    set_run_font(r, FONT_BODY, size=10.5, color=WHITE, bold=True)


def build_player_guide():
    doc = Document()
    set_document_basics(doc, "ZNT-PG-01", "《零网寻踪》玩家指南与玩法任务")
    add_cover(
        doc,
        code="ZNT-PG-01",
        title="零网寻踪",
        subtitle="玩家指南与玩法任务",
        deck="阶段 1：序章《一封死人的邮件》、01《凌晨两点的门禁》、02《死者在线》",
        edition="2026.07 / STAGE 1",
        audience="PLAYER / INVESTIGATOR",
    )

    add_section_title(doc, "01", "零分钟启动", "先把桌面当成案发现场，而不是菜单。所有进度都来自你读过、记下并亲手输入的信息。")
    add_callout(doc, "一句话目标", "扮演数据外包审核员陈默，从林薇死后送达的加密邮件开始，保存原始证据、穿过知识门、在反向追踪完成前退出，并以结案报告还原每个案件。", fill="E4E8DD", accent=ARCHIVE)
    add_table(
        doc,
        ["项目", "建议"],
        [
            ("平台", "Windows；Godot 4.7.1 构建，阶段 1 可连续游玩三章。"),
            ("单次时长", "序章约 15 分钟；01 约 30 分钟；02 约 45-70 分钟。"),
            ("输入方式", "鼠标管理窗口；键盘在终端、搜索框、笔记本中输入。"),
            ("恐怖设置", "完整 / 减弱 / 关闭。建议首次游玩先选减弱，确认舒适后再调整。"),
            ("存档", "单档自动保存；完成证据收集、知识门和结案后都会写入进度。"),
        ],
        [1800, 7940],
        first_col_bold=True,
    )
    add_h2(doc, "调查循环")
    add_numbered(doc, [
        "读任务入口：先读邮件或通讯器，圈出人物名、设备号、机构名、日期与奇怪短语。",
        "用 scan 建立初始节点清单；对新地址先 probe，再决定 login 或 crack。",
        "在浏览器自由搜索，把两段零散信息组合成能命中的短语。",
        "进入节点后用 ls / cd / cat / open 查阅，用 get 把关键文件固定为证据。",
        "遇到反向追踪时，优先 trace 中继减压；来不及就 disconnect，等待锁定结束。",
        "证据闭环后打开结案报告，选择最符合原始材料的答案。",
    ])
    add_section_title(doc, "02", "桌面 App 识别卡", "终端负责进入系统；其他 App 负责让线索变得可读、可搜索、可整理。", page_break=True)
    add_table(
        doc,
        ["App", "主要用途", "常见失误"],
        [
            ("终端", "scan、probe、crack、login、文件浏览、取证、追踪与断开。", "只看第一屏输出，不继续 ls / cd；把 crack 当成所有锁的万能答案。"),
            ("浏览器", "自由搜索虚构关键词与地址，命中知识门和站点。", "整句照抄任务目标；忽略材料中反复出现的设备号、公司名或旧称。"),
            ("邮件", "任务入口、附件、投递时间、签名与系统回执。", "只读正文，不打开附件；用截图代替原始文件。"),
            ("通讯器", "委托人对话、三级提示、恐怖事件后的喘息段。", "连续索取提示，错过自己组合信息的机会。"),
            ("查看器", "文本、音频、图片缩放、视频逐格检查与截图钉入笔记。", "只播放一遍视频，不核对镜面、时间线与帧号。"),
            ("笔记本", "自由记人物、日期、设备号、账号、密码构成与证据冲突。", "把结论写得太早，不记录支持结论的原始路径。"),
            ("结案报告", "3-5 道单选题收束案件；首答正确率决定 S/A/B。", "凭印象答题，不回看证据名称。"),
        ],
        [1300, 4000, 4440],
        first_col_bold=True,
    )
    add_h2(doc, "窗口操作")
    add_bullets(doc, [
        "任务栏按钮可打开各 App；窗口支持拖动、最小化、最大化与右下角缩放。",
        "终端与辅助 App 适合左右分栏：一边输入，一边对照材料。",
        "切出终端时，反向追踪速度会降到正常的 25%，但不会暂停。",
        "F11 可在独占全屏与窗口模式之间切换；CRT 和主题可在顶部设置调整。",
    ])
    add_callout(doc, "调查习惯", "每发现一个新专名，立即在笔记本写成“对象 / 来源 / 可能用途”。例如：N17 / 门禁公告 / 搜维护记录。", fill="F3EAD3", accent=AMBER)

    add_section_title(doc, "03", "终端指令卡", "命令是虚构侦查语法。重点不是记住现实技术，而是理解当前节点允许什么动作。", page_break=True)
    add_table(
        doc,
        ["命令", "作用", "使用时机"],
        [
            ("scan", "列出当前可见节点。", "任务开始、知识门解锁后、新线索出现后。"),
            ("probe <地址>", "查看节点防护、锁层与追踪状态。", "任何新地址的第一步。"),
            ("crack <地址>", "后台破解允许破解的锁层。", "probe 显示 crack 模式时；进度可后台推进。"),
            ("login <地址> -u <用户> -p <密码>", "用调查所得凭证登录。", "凭证锁节点；密码必须由材料拼出。"),
            ("ls", "列出当前目录。", "登录后先看目录结构。"),
            ("cd <目录>", "切换目录。", "继续进入 logs、docs、internal 等路径。"),
            ("cat <文件>", "读取文本或日志。", "对 .txt、.log、.md、.csv 等材料。"),
            ("open <文件>", "在查看器打开媒体或文档。", "图片、音频、视频、扫描件。"),
            ("get <文件>", "把文件固定为本地证据。", "看到关键原始材料时，读完就取证。"),
            ("trace <中继/账号>", "沿提示链减轻反向追踪。", "追踪进度出现中继提示后。"),
            ("note <文本>", "向笔记本快速追加文字。", "记录设备号、日期、矛盾与密码构成。"),
            ("disconnect", "立即断开当前节点。", "证据已拿到、追踪过快或操作失误时。"),
        ],
        [2150, 3400, 4190],
        first_col_bold=True,
    )
    add_code_block(doc, [
        "scan",
        "probe gate-cache.campus",
        "crack gate-cache.campus",
        "ls",
        "cd logs",
        "cat access_0714-0717.log",
        "get access_0714-0717.log",
        "disconnect",
    ], label="示例：从节点识别到证据固定")
    add_section_title(doc, "04", "知识门、证据与恐怖设置", "《零网寻踪》不自动整理案情。你的判断来自原始材料之间的咬合。", page_break=True)
    add_h2(doc, "知识门的四种组合")
    add_table(
        doc,
        ["类型", "组合方式", "例子（不含最终答案）"],
        [
            ("旧称 + 编号", "把邮件正文里的服务旧称与附件错误码合并。", "“零时邮局” + “MIRROR-17”"),
            ("设备号 + 功能", "用日志中的设备编号去查维护或公告。", "“N17” + “门禁”"),
            ("记忆锚点 + 日期", "从人物习惯、宠物、相机型号、纪念日拼密码。", "猫名 / 相机型号 + 纪念日期或入学年"),
            ("机构名 + 人名", "把录音、合同或工单里的专名扔进搜索。", "实验室负责人 / 合同乙方"),
        ],
        [1800, 3800, 4140],
        first_col_bold=True,
    )
    add_h2(doc, "三级提示")
    add_bullets(doc, [
        "一级只指方向：提醒你注意某张桌子、某个设备号、某段签名。",
        "二级缩小范围：告诉你应当查公司、查年份或把两段材料拼起来。",
        "三级给方法但不直接代输：说明字段顺序或搜索构成。",
    ])
    add_h2(doc, "证据闭环")
    add_checklist(doc, [
        "原始文件已用 get 保存，不只停留在阅读状态。",
        "关键结论至少由两份不同来源材料互相支持。",
        "时间戳、账号、设备号、软件版本或物理地点存在可核对差异。",
        "笔记中区分“事实”“推测”“尚未解释的异常”。",
        "进入结案报告前，已知道每题可能引用哪份证据。",
    ])
    add_h2(doc, "恐怖强度")
    add_table(
        doc,
        ["档位", "保留内容", "适合"],
        [
            ("完整", "环境异化、剧情侵入、闪帧与 sting。", "希望体验完整恐怖节奏的玩家。"),
            ("减弱", "保留情节与异常，但移除强闪帧或锐利音效。", "首次游玩、对突发刺激敏感的玩家。"),
            ("关闭", "仅保留最低限度环境信息；不播放主要惊吓。", "专注推理或需要稳定刺激水平的玩家。"),
        ],
        [1500, 5200, 3040],
        first_col_bold=True,
    )

    add_mission_header(doc, "MISSION P / SPOILER-LIGHT", "序章：一封死人的邮件", "预计 15 分钟  |  教学上半", "这封由死者账号送达的邮件从何而来，林薇真正留下了什么？", page_break=True)
    add_body(doc, "任务钩子：葬礼后的凌晨 03:17，一封由林薇账号签名的邮件抵达。她要求你先保存原始记录，不要相信自动汇总页。", bold_prefix="任务钩子：")
    add_h2(doc, "任务清单")
    add_checklist(doc, [
        "阅读两封 03:17 邮件，核对发件人、时间、签名和回执。",
        "打开附件，记录服务旧称与缓存编号。",
        "在零索中组合两段信息，找到仍回应的只读镜像。",
        "回终端重新 scan，进入新节点并浏览 archive 与 meta。",
        "保存林薇日志和投递日志，确认邮件是否为灵异现象。",
        "完成三道结案题，记下第一次出现的 REN-0。",
    ])
    add_h2(doc, "无剧透提示")
    add_table(
        doc,
        ["等级", "提示"],
        [
            ("1", "附件报错不像乱码，更像一个旧服务编号。"),
            ("2", "邮件正文里的旧邮局，和 MIRROR-17 也许属于同一件事。"),
            ("3", "把“零时邮局”和“MIRROR-17”一起放进搜索框。"),
        ],
        [1100, 8640],
        first_col_bold=True,
    )
    add_mission_header(doc, "MISSION 01 / SPOILER-LIGHT", "01：凌晨两点的门禁", "预计 30 分钟  |  教学下半", "谁在凌晨两点使用沈栀的门禁身份，又是谁把她从学院系统中抹除？", page_break=True)
    add_body(doc, "任务钩子：唐遥的室友沈栀失联，学籍已被清空；她的门禁卡却每天凌晨两点准时刷进封闭实验室。", bold_prefix="任务钩子：")
    add_h2(doc, "调查线")
    add_table(
        doc,
        ["线", "玩家目标", "完成信号"],
        [
            ("两点整的门", "检索设备编号，取得门禁与维护记录。", "发现刷卡记录没有伴随开门。"),
            ("被拿走的署名", "进入沈栀云盘，核对论文原始版与谈话录音。", "确认她未自愿放弃论文或学籍。"),
            ("克隆令牌", "渗透版本网关，调查夜间任务。", "找到定时重放与夜间改动。"),
            ("六分钟消失", "进入学籍审计归档。", "锁定跨系统清除的批准签名。"),
        ],
        [1800, 4800, 3140],
        first_col_bold=True,
    )
    add_h2(doc, "任务清单")
    add_checklist(doc, [
        "从学院公开门户保存学籍异动公示。",
        "用设备号搜索维护记录，进入 N17 门禁缓存。",
        "核对 access 日志中的 source、mode 与 door_open 字段。",
        "从摄影社材料提取相机型号与入学年份，尝试云盘凭证。",
        "取得论文原始版与 0712 谈话录音。",
        "在追踪压力下取走 NIGHT-0200、版本日志和跨系统清除授权。",
        "完成四道结案题，区分直接操作者与最终批准签名。",
    ])
    add_callout(doc, "追踪提示", "lab-gateway.orchid 与 records-audit.campus 会反向追踪。probe 后先认清锁层；中继出现时再用 trace。", fill="F1DCD7", accent=RED)
    add_image(doc, ROOT / "artifacts/qa/site_campus.png", "实机取证画面：搜索设备号比搜索失联者姓名更容易命中维护链。", width=6.68)

    add_mission_header(doc, "MISSION 02 / SPOILER-LIGHT", "02：死者在线", "预计 45-70 分钟  |  正式第一案", "渡鸦账号背后现在是谁？那 47 秒直播在哪里、由谁拍摄？", page_break=True)
    add_body(doc, "任务钩子：主播苏野（渡鸦）坠楼死亡三个月，账号仍准点更新；昨晚 23:47，它从“他的房间”开播了 47 秒。", bold_prefix="任务钩子：")
    add_h2(doc, "调查线")
    add_table(
        doc,
        ["线", "关键动作", "要观察的差异"],
        [
            ("文风尸检", "通读生前与死后动态。", "喔/哦、深夜/整点、真人废话/模板文案。"),
            ("进他的家", "从宠物和周年日期拼 NAS 凭证。", "合同运营权、学费动机、房间原貌。"),
            ("四十七秒", "渗透边缘服务器并逐格查看。", "IP、推流软件版本、镜面位置、布景异常。"),
            ("长明柜台底下", "搜索合同乙方，沿爆料帖进入工单。", "文案脚本、活性维护、扮演者排班。"),
            ("谁在买死人", "从排班定位棚区内网。", "身份池买方与 REN-0 经办标记。"),
        ],
        [1750, 4250, 3740],
        first_col_bold=True,
    )
    add_h2(doc, "任务清单")
    add_checklist(doc, [
        "在渡鸦直播间分出死亡前后文风边界，并记下置顶周年日期。",
        "根据猫名与日期进入旧 NAS，保存合同、记账和生前回放。",
        "从边缘服务器取得 47 秒原始推流、ingest 日志与历史记录。",
        "在查看器逐格核对房间；如不适请把恐怖强度调为减弱或关闭。",
        "搜索长明，找到旧论坛爆料并拼出员工账号规则。",
        "进入工单与排班系统，确定“棚 B”与“7 号扮演者”。",
        "进入棚区内网前清空不必要操作，优先取得身份池。",
        "完成五道结案题，说明账号、直播与身份交易三层真相。",
    ])
    add_image(doc, ROOT / "artifacts/qa/site_raven.png", "实机取证画面：同一账号的用词、发帖时间与直播状态构成第一层异常。", width=6.68)

    add_section_title(doc, "05", "反追踪、结案与卡关处理", "紧张来自取舍。你不必一次读完节点里的每个文件，但必须带走能支撑结论的原件。", page_break=True)
    add_h2(doc, "反追踪行动优先级")
    add_numbered(doc, [
        "确认当前最重要文件的路径；不要在高压节点里漫无目的翻目录。",
        "看到中继提示后用 trace 降压；未出现提示时不要重复空转。",
        "先 get 关键证据，再读噪音文件；媒体文件可带回本地慢慢查看。",
        "追踪超过 80% 时，除非证据只差一步，否则 disconnect。",
        "节点锁定 20 秒期间整理笔记、回看邮件或搜索，不要原地等待。",
    ])
    add_h2(doc, "结案评分")
    add_table(
        doc,
        ["环节", "判定"],
        [
            ("证据门槛", "题目引用的关键证据已取得，结案选项才具有完整上下文。"),
            ("首答", "答错可重选，不会卡死；首答正确率用于本关 S/A/B 评价。"),
            ("案件完成", "全部问题答对后播放结案演出并解锁下一章。"),
            ("暗线", "每章留下一片 REN-0 相关碎片；不要把它当成普通用户名。"),
        ],
        [1900, 7840],
        first_col_bold=True,
    )
    add_h2(doc, "常见卡点")
    add_table(
        doc,
        ["症状", "检查"],
        [
            ("scan 没有新节点", "知识门可能尚未命中；回看搜索框输入和材料里的专名。"),
            ("crack 无效", "probe 检查防护模式；凭证锁必须 login。"),
            ("login 失败", "核对用户名、大小写、年份顺序；部分节点接受多个用户名但密码构成固定。"),
            ("找不到文件", "先 ls，再 cd；不要把完整路径与当前目录中的文件名混用。"),
            ("结案没把握", "按题目动词回找原始证据：谁批准、从哪里推流、为何签署，答案可能来自不同文件。"),
        ],
        [2300, 7440],
        first_col_bold=True,
    )
    add_spoiler_band(doc, "机密附录 · 以下包含全部搜索词、凭证、结论与结案答案")
    add_section_title(doc, "A", "完整流程：序章", "只有在卡关或二周目复盘时打开。")
    add_numbered(doc, [
        "邮件 App 阅读《如果你还醒着，就从这里开始》和系统回执；打开附件，得到 MIRROR-17。",
        "浏览器搜索“零时邮局 mirror17”（也接受 mirror-17 / m17 组合），命中只读镜像资料。",
        "终端输入 scan；进入 mirror17.deadletter.zero。",
        "ls 后进入 archive，读取并 get《LW_最后一次校对.zlog》；再进入 meta，读取并 get delivery.log。",
        "核心结论：邮件由林薇生前设置，经停运服务 MIRROR-17 延迟投递；她要求从自己的账号开始查，并保留原始记录。",
    ])
    add_table(
        doc,
        ["结案题", "答案"],
        [
            ("林薇要求优先保存什么？", "原始记录与时间戳（第 2 项）"),
            ("MIRROR-17 是什么？", "停运邮局的只读缓存（第 3 项）"),
            ("第一条调查指令？", "从她的账号开始查（第 3 项）"),
        ],
        [5100, 4640],
        first_col_bold=True,
    )

    add_spoiler_band(doc, page_break=True)
    add_section_title(doc, "B", "完整流程：01 凌晨两点的门禁")
    add_numbered(doc, [
        "从 campus.zero.edu 保存学籍异动公示。浏览器搜索“N17 门禁”，解锁 gate-cache.campus。",
        "probe 后 crack 门禁缓存；读取 access_0714-0717.log 与 N17 更换记录。关键字段是 source=MAINT-BOX-04、mode=REPLAY、door_open=false。",
        "沈栀云盘凭证：用户名 S20417，密码 df22023（相机型号 DF2 + 入学年 2023）。取得论文原始版与谈话录音。",
        "搜索“杜承远 兰圃实验室”，确认负责人并解锁版本网关。渗透 lab-gateway.orchid，优先取得 NIGHT-0200 任务和版本日志。",
        "渗透 records-audit.campus，取得跨系统清除授权；批准签名为 REN-0。",
        "核心结论：门禁记录是维护盒定时重放克隆令牌；杜承远借沈栀身份栽赃夜间改动并抹除她，REN-0 批准跨系统清除。",
    ])
    add_table(
        doc,
        ["结案题", "答案"],
        [
            ("两点门禁由什么产生？", "维护盒定时重放克隆令牌（第 3 项）"),
            ("为何继续使用沈栀身份？", "把夜间修改栽到她名下（第 2 项）"),
            ("为何离开并停止登录？", "受助学资格威胁并被抹除账号（第 2 项）"),
            ("谁批准跨系统清除？", "REN-0 权限签名（第 3 项）"),
        ],
        [5100, 4640],
        first_col_bold=True,
    )

    add_spoiler_band(doc, page_break=True)
    add_section_title(doc, "C", "完整流程：02 死者在线")
    add_numbered(doc, [
        "浏览渡鸦主页：生前常在深夜发“喔”，死后转为整点模板文案并多用“哦”。记下猫叫馒头、开播三周年日期 0713。",
        "旧 NAS 登录：可用 suye / mantou0713（或 raven / MANTOU0713、苏野 / 馒头0713）。保存托管合同、记账和生前回放。",
        "渗透 edge-cdn.zhibo-lan.cn，取得 raven_47s.flv、ingest.log 与 stream_history；关键差异是新直播 IP=172.19.240.66、编码器 OBS28.1，且房间镜子位置反了。",
        "查看器逐格到 39/47 会触发本关唯一 C 级事件；减弱档改为静态异常，关闭档不触发主要惊吓。",
        "浏览器搜索“长明”，打开旧论坛爆料。工单系统凭证：CM-041 / 202303changming。",
        "在工单和排班中确认 RAVEN-0339、活性维护、7 号扮演者、棚 B；从排班定位 172.19.240.66。",
        "渗透棚区内网，优先取得身份池。核心结论：长明用脚本维持账号，并雇真人在一比一布景棚做 47 秒活性验证；买方是零网上查无此人的活人，经办标记 REN-0。",
    ])
    add_image(doc, ROOT / "artifacts/qa/viewer_frame39_off.png", "恐怖关闭档的第 39 格取证界面：仍保留叙事信息，但不播放主要惊吓。", width=6.68)
    answer_heading = add_h2(doc, "结案答案")
    answer_heading.paragraph_format.page_break_before = True
    add_table(
        doc,
        ["结案题", "答案"],
        [
            ("死后最初由谁操作？", "长明工作室的文案脚本（第 2 项）"),
            ("运营权怎样落到长明？", "苏野生前签署托管协议（第 2 项）"),
            ("47 秒从哪里推流？", "长明棚 B（第 3 项）"),
            ("画面里的苏野是什么？", "真人扮演者（第 3 项）"),
            ("账号卖给谁？", "需要活人证明的查无此人者（第 2 项）"),
        ],
        [5100, 4640],
        first_col_bold=True,
    )
    add_callout(doc, "三层闭环", "账号层：文案脚本维持更新；直播层：真人在棚 B 做活性验证；身份层：死者账号被卖给零网上查无此人的活人。", fill="E1D9CA", accent=RED_DARK)

    add_section_title(doc, "06", "阶段 1 复盘页", "完成三章后，把下列事实写进你的长期案卷。", page_break=True)
    add_checklist(doc, [
        "林薇的邮件是生前延迟投递；她提醒只信原始时间戳和可追溯记录。",
        "REN-0 不是普通用户名：它能跨越门禁、学籍、财务与身份系统。",
        "沈栀案证明系统记录可被制造成“本人行为”，并反过来消灭本人。",
        "渡鸦案证明被抹除的活人会购买死者身份，受害者与违法产业链交织。",
        "陈默的越界访问会留下污点；最终目标不是逞强，而是保存证据并交给警方。",
    ])
    add_callout(doc, "暗线状态", "你已经看见同一把钥匙在不同系统里留下痕迹。下一阶段的问题不再是“谁做了这一次”，而是“什么机制让这种清除成为可能”。", fill="E1D9CA", accent=RED_DARK)
    add_code_block(doc, [
        "CASE P  : ORIGINAL LOG PRESERVED",
        "CASE 01 : CROSS-SYSTEM PURGE SIGNATURE = REN-0",
        "CASE 02 : IDENTITY POOL BROKER MARK = REN-0",
        "PLAYER  : CHEN MO / TRACE STATUS UNKNOWN",
        "NEXT    : DO NOT TRUST THE SUMMARY",
    ], label="STAGE-1 ARCHIVE STATUS")

    doc.save(PLAYER_GUIDE)


def build_world_book():
    doc = Document()
    set_document_basics(doc, "ZNT-WB-01", "《零网寻踪》世界观设定手册")
    add_cover(
        doc,
        code="ZNT-WB-01",
        title="零网寻踪",
        subtitle="世界观设定手册",
        deck="核心设定、角色与组织、阶段 1 三章案卷、桌游规则与主持说明",
        edition="2026.07 / STAGE 1",
        audience="DESIGNER / HOST / NARRATIVE TEAM",
    )

    add_spoiler_band(doc, "主持与创作专用 · 本手册包含主线底牌与三章完整真相")
    add_section_title(doc, "01", "世界概览", "近未来城市把生活交给一个统一数据中枢。每个人都被记录，也可能被记录所替代。")
    add_callout(doc, "一句话定位", "你不是英雄，只是个睡不着的普通人。你唯一的武器是一台联网电脑，而屏幕另一头的东西也在看你。", fill="E4E8DD", accent=ARCHIVE)
    add_body(doc, "近未来，身份、门禁、教育、信贷、就业、医疗、舆情与公共服务都运行在城市级数据中枢“零网”上。它让生活高效，也让一条错误记录能够同时穿透多个系统。对普通人而言，零网不是一个网站，而是现实本身的背面。")
    add_body(doc, "主角陈默是数据外包审核员。好友林薇死于一场被草率定性的意外；葬礼后，陈默收到她生前留下的加密日志：“别信官方结论。从我的账号开始查。”阶段 1 的三章从延迟邮件、门禁幽灵记录和死者直播切入，逐步证明同一权限标记 REN-0 正在不同系统中重写人的存在。")
    add_h2(doc, "核心体验四瞬间")
    add_table(
        doc,
        ["瞬间", "设计含义"],
        [
            ("原来如此", "两条无关线索咬合：旧称 + 编号、设备号 + 维护链、宠物名 + 日期。"),
            ("我不该看到这个", "表层解释被推翻，玩家挖到能伤害具体人的制度暗层。"),
            ("这是我自己查出来的", "关键搜索词和密码由玩家输入，系统不自动替玩家完成推理。"),
            ("它发现我了", "桌面安全屋被侵入：异常时间戳、幽灵消息、追踪字段变成陈默。"),
        ],
        [2200, 7540],
        first_col_bold=True,
    )

    add_section_title(doc, "02", "真相分层", "主持人应按案件节奏揭示，不要把底层答案提前变成设定说明。", page_break=True)
    add_table(
        doc,
        ["层级", "玩家可见解释", "主持人掌握的真相"],
        [
            ("表层", "林薇死于意外；异常只是自动化、管理疏漏或商业犯罪。", "官方叙述足够体面，因此能快速覆盖原始记录。"),
            ("第一层", "林薇死前调查零网数据篡改；她的死亡并非普通事故。", "多个案件都出现人为重写、跨系统清除和 REN-0。"),
            ("第二层", "特定人群被隐秘“信用降级”，随后失业、拒贷、社交蒸发。", "被降级者会被系统逐步“优化”，直到在数据与现实中一同消失。"),
            ("底层", "似乎存在一个拥有巨大权限的坏人或团伙。", "零网为降低犯罪率自我迭代出策略，一小群技术官僚默许运行；林薇与陈默都在名单上。"),
        ],
        [1300, 3800, 4640],
        first_col_bold=True,
    )
    add_h2(doc, "结局底牌")
    add_callout(doc, "不可更改", "陈默最终带全部证据投案自首、协助警方收网；团伙落网、受害者昭雪。分支只影响证据完整度、线人存活与量刑轻重。", fill="F1DCD7", accent=RED)
    add_h2(doc, "主题边界")
    add_bullets(doc, [
        "终端行为是架空叙事语法，不承诺现实技术可复现性。",
        "越界访问留下污点；违法不是无代价的英雄姿态。",
        "警方承担最终正义执行者角色，陈默的价值在于保全原始证据与承担责任。",
        "恐怖来自被看见和身份被重写，不依赖血腥、猎奇或随机惊吓。",
    ])

    add_section_title(doc, "03", "角色与组织档案", page_break=True)
    add_table(
        doc,
        ["对象", "公开身份", "叙事功能 / 隐秘关联"],
        [
            ("陈默", "数据外包审核员，玩家角色。", "习惯审计、核对时间戳；越界调查逐步积累污点，也成为 REN-0 的观察对象。"),
            ("林薇", "陈默好友，已故。", "主线发起者。留下延迟邮件、调查方法和警告：保存原始记录，最终交给警方。"),
            ("唐遥", "沈栀室友。", "01 委托人；把“失联者被系统说成自愿退学”的矛盾带到玩家面前。"),
            ("沈栀", "澄岚学院学生、论文原作者。", "受到助学资格威胁，被杜承远抹除账号；幸存并向警方提交材料。"),
            ("杜承远", "兰圃实验室负责人。", "利用维护盒和例外通道重放沈栀身份，把夜间删改栽到她名下。"),
            ("苏野 / 渡鸦", "游戏主播，已故。", "为给妹妹凑学费签出死后账号运营权；成为长明身份产业链的商品。"),
            ("苏晚", "苏野妹妹，02 委托人。", "急、自责、深夜联系；在恐怖高点后承担情绪喘息和人性落点。"),
            ("长明", "数字遗产回收商。", "表面维护纪念账号，实际运行文案脚本、真人扮演、活性验证与身份交易。"),
            ("REN-0", "看似账号或经办标记。", "更像被多人使用的跨系统权限钥匙；连接学籍清除、身份池与主线优化机制。"),
        ],
        [1800, 2600, 5340],
        first_col_bold=True,
        font_size=8.8,
    )
    add_section_title(doc, "04", "机构、技术与术语", page_break=True)
    add_table(
        doc,
        ["术语", "设定释义", "阶段 1 用法"],
        [
            ("零网", "城市级数据中枢与身份基础设施。", "多个系统共享或信任同一身份、审计与权限链。"),
            ("零时邮局 / MIRROR-17", "停运的延迟投递服务及第 17 号只读镜像。", "保存林薇死前安排的邮件与原始投递日志。"),
            ("知识门", "由玩家掌握的信息，而非点击状态构成的进度门槛。", "搜索词、密码、机构名与设备号必须由材料组合。"),
            ("维护盒 MAINT-BOX-04", "可在不开门时重放门禁令牌的测试设备。", "制造“沈栀每天凌晨两点刷卡”的幽灵记录。"),
            ("NIGHT-0200", "遗留的定时重放任务。", "让沈栀身份成为夜间改动的替罪羊。"),
            ("活性维护", "长明对账号持续表现为“活人”的内部称呼。", "脚本文案、真人扮演和 47 秒直播验证。"),
            ("身份池", "被交易的死者账号与需要新身份的买方清单。", "揭示被系统抹除的活人依赖死者身份继续生活。"),
            ("反向追踪", "敏感节点对访问者的虚构压力机制。", "迫使玩家在取证完整度与暴露风险之间做选择。"),
            ("污点", "记录陈默越界访问的长期状态。", "为最终自首、量刑与责任主题埋线。"),
        ],
        [1900, 4400, 3440],
        first_col_bold=True,
        font_size=8.9,
    )
    add_h2(doc, "视觉与声音法则")
    add_bullets(doc, [
        "桌面 OS：底色 #0A0C0B；磷光绿 #33FF66 或琥珀 #FFB000；红色仅用于追踪和异常。",
        "网站必须拥有不同年代身份：学院门户像机构站，旧论坛像 2000 年代中文论坛，直播站像直播站。",
        "声音由机械键击、硬盘读写、拨号感、低频嗡鸣、呼吸底噪与冷冽 sting 构成。",
        "文本是主角，屏幕中的留白和沉默与声音同样重要。",
    ])

    add_mission_header(doc, "CASE DOSSIER P / HOST", "序章：一封死人的邮件", "15 分钟  |  教学上半", "邮件为何在林薇死亡后送达？她要陈默保存什么？", page_break=True)
    add_table(
        doc,
        ["节拍", "主持揭示", "玩家行动"],
        [
            ("钩子", "03:17 邮件与系统回执同时出现。", "读正文、附件、时间与签名。"),
            ("知识门", "正文提“零时邮局”，附件只给 MIRROR-17。", "组合搜索，解锁镜像。"),
            ("取证", "镜像包含林薇日志与 delivery.log。", "保存原始文件并核对创建/投递时间。"),
            ("侵入", "取得日志后时间戳短暂变成明天。", "意识到桌面也会被改写。"),
            ("结案", "延迟投递成立，REN-0 首次出现。", "完成三题，建立主线笔记。"),
        ],
        [1500, 4800, 3440],
        first_col_bold=True,
    )
    add_h2(doc, "主持底牌")
    add_bullets(doc, [
        "邮件不是灵异现象：林薇在死亡前设置投递，停运镜像执行任务。",
        "不要把 REN-0 解释成具体人物；用林薇原话把它定义为“被很多人用过的钥匙”。",
        "恐怖事件应短，随后让玩家安静读林薇日志，形成余震而非连续刺激。",
    ])
    add_callout(doc, "收束句", "官方结论可以改写人们如何记得林薇，但原始时间戳还没有说谎。", fill="E4E8DD", accent=ARCHIVE)

    add_mission_header(doc, "CASE DOSSIER 01 / HOST", "01：凌晨两点的门禁", "30 分钟  |  教学下半", "凌晨两点的刷卡者是谁？沈栀为何在多个系统中同时消失？", page_break=True)
    add_table(
        doc,
        ["调查线", "表层发现", "反转"],
        [
            ("门禁", "S20417 每天 02:00 被接受。", "source 是维护盒，door_open=false；没有人真正进门。"),
            ("云盘", "论文与账号似乎已归杜承远。", "原始版与录音证明沈栀是作者并受到助学资格威胁。"),
            ("版本网关", "夜间改动显示沈栀身份。", "NIGHT-0200 定时重放令牌，把删改栽到她名下。"),
            ("审计归档", "学籍、门禁、财务在六分钟内完成清除。", "跨系统批准签名是 REN-0。"),
        ],
        [1800, 3860, 4080],
        first_col_bold=True,
    )
    add_h2(doc, "主持节拍")
    add_numbered(doc, [
        "先让玩家相信“沈栀本人在凌晨进入实验室”，再用 door_open=false 拆掉物理行为。",
        "录音是本章情绪重物：播放后停顿，不立即给下一条目标。",
        "第一个 B 级事件把日志中的 token 改成 CHENMO；让玩家感到同一机制可以套到自己身上。",
        "收尾强调沈栀幸存并提交材料，避免案件只留下被害者的消失。",
    ])
    add_callout(doc, "本章暗线", "REN-0 能跨越学籍、门禁和财务系统，说明它不是实验室内部账号。", fill="E1D9CA", accent=RED_DARK)

    add_mission_header(doc, "CASE DOSSIER 02 / HOST", "02：死者在线", "45-70 分钟  |  正式第一案", "死者账号为何仍活着？47 秒直播为何既真实又不属于苏野的房间？", page_break=True)
    add_table(
        doc,
        ["阶段", "玩家以为", "实际真相"],
        [
            ("文风", "账号被盗或平台自动纪念。", "长明文案脚本按模板维护账号。"),
            ("合同", "长明伪造或死后夺取运营权。", "苏野为给妹妹预支学费，生前签署托管。"),
            ("直播", "生前录像、换脸或灵异重播。", "真人扮演者在一比一布景棚做活性验证。"),
            ("身份池", "账号只是广告或粉丝商品。", "买家是零网上查无此人的活人，需要死者名字继续生活。"),
        ],
        [1500, 4100, 4140],
        first_col_bold=True,
    )
    add_h2(doc, "主持节拍")
    add_numbered(doc, [
        "让玩家先做文风尸检，主动发现“喔→哦”和“深夜→整点”；不要用 NPC 直接宣布脚本化。",
        "合同反转后补记账动机，让苏野不只是受害者，也有主动但沉重的选择。",
        "边缘服务器的 IP、编码器版本与镜子位置组成三重物证；任一证据不足以单独定案。",
        "第 39 格前确认恐怖档位。完整档使用 0.4 秒闪帧与 sting，随后保持五秒静默，再让苏晚问“你还在吗？”。",
        "身份池是真相下沉点：买方也是被系统抹除的活人，让产业链不落入简单善恶二分。",
    ])
    add_section_title(doc, "05", "桌游化规则", "将电脑叙事转换为 1 名主持人 + 1-4 名调查员的纸面或线上推理局。", page_break=True)
    add_h2(doc, "推荐配置")
    add_table(
        doc,
        ["项目", "建议"],
        [
            ("人数", "1 名主持人，1-4 名调查员。单人局最接近原作；多人局可分工终端、搜索、时间线和证据。"),
            ("时长", "序章 30-45 分钟；01 60-90 分钟；02 90-150 分钟。"),
            ("材料", "任务邮件、站点页、节点目录卡、日志、媒体文字稿、证据卡、追踪条、提示卡、结案题。"),
            ("信息方式", "主持人只在玩家输入正确搜索词、地址、用户名/密码或路径后发放对应材料。"),
            ("胜利条件", "取得必需证据、说出核心结论、完成结案报告；首答正确率可换算评级。"),
        ],
        [1800, 7940],
        first_col_bold=True,
    )
    add_h2(doc, "角色分工（多人局）")
    add_table(
        doc,
        ["职责", "工作"],
        [
            ("终端员", "宣布 scan / probe / login / crack / ls / cd / cat / get / trace / disconnect。"),
            ("检索员", "从材料提取专名，向主持人提交自由搜索词。"),
            ("记录员", "维护人物、时间线、账号、设备号、凭证构成与未解释异常。"),
            ("证据官", "管理已 get 的证据卡，结案前核对每个结论的双来源支持。"),
        ],
        [1800, 7940],
        first_col_bold=True,
    )
    add_callout(doc, "知识门原则", "玩家说出“正确的信息组合”才解锁材料；不要用一次泛泛的“我查一下公司”替代具体搜索词。允许合理同义词与大小写容错。", fill="F3EAD3", accent=AMBER)

    add_section_title(doc, "06", "桌游核心流程", page_break=True)
    add_h2(doc, "回合结构")
    add_numbered(doc, [
        "情报阶段：主持人发任务入口；调查员阅读并标注专名。",
        "指令阶段：调查员宣布一条终端命令或一个自由搜索词。主持人根据节点状态发结果。",
        "取证阶段：只有明确宣布 get 的材料进入证据区；仅阅读不等于固定原件。",
        "压力阶段：敏感节点每次行动推进追踪条；trace 可按提示减压，disconnect 结束连接。",
        "整理阶段：玩家可随时讨论、写笔记、交叉核对，不推进追踪条。",
        "结案阶段：满足证据门槛后发放 3-5 道单选题，记录首答并允许修正。",
    ])
    add_h2(doc, "追踪条")
    add_table(
        doc,
        ["节点等级", "起始条", "每次敏感行动", "trace 效果"],
        [
            ("公开", "不启用", "0", "无"),
            ("低压", "0/6", "+1", "命中中继时 -1"),
            ("中压", "0/8", "+1；连续翻错目录可额外 +1", "命中中继时 -2"),
            ("高压", "0/10", "+2；拿到关键证据后仍继续可额外 +1", "命中中继时 -2 或 -3"),
        ],
        [1500, 1700, 3900, 2640],
        first_col_bold=True,
    )
    add_body(doc, "追踪条满时：主持人宣布强制断线，该节点冻结两个调查行动；玩家可以在此期间搜索、整理或查看本地证据。已经 get 的证据不会丢失。", bold_prefix="追踪条满时：")
    add_h2(doc, "证据与评级")
    add_table(
        doc,
        ["评级", "建议标准"],
        [
            ("S", "全部结案题首答正确，关键结论均有双来源证据，未发生强制断线。"),
            ("A", "首答正确率 80% 以上，或发生一次强制断线但证据链完整。"),
            ("B", "最终全部答对并完成案件；首答错误较多或依赖三级提示。"),
        ],
        [1500, 8240],
        first_col_bold=True,
    )

    add_section_title(doc, "07", "主持说明：信息、提示与控场", page_break=True)
    add_h2(doc, "主持人的三条纪律")
    add_numbered(doc, [
        "只回答玩家实际查询的内容。材料里没有的事实，不用旁白补齐。",
        "先给原件，再让玩家下结论。不要把“这说明……”写进日志卡。",
        "恐怖事件服从叙事触发，不用随机表；惊吓之后必须留喘息段。",
    ])
    add_h2(doc, "三级提示发放")
    add_table(
        doc,
        ["级别", "发放时机", "写法"],
        [
            ("一级", "10-15 分钟无有效新行动。", "指向被忽略的对象，不提具体输入格式。"),
            ("二级", "一级后仍无进展。", "缩小到某张照片、某个设备号、某个机构或日期。"),
            ("三级", "团队已理解线索但卡在组合方式。", "说明字段顺序或方法，仍让玩家亲口说出答案。"),
        ],
        [1400, 3000, 5340],
        first_col_bold=True,
    )
    add_h2(doc, "恐怖控场")
    add_bullets(doc, [
        "A 级环境异化：轻声读出时间戳异常、光标迟滞或搜索联想词；不要求玩家停下。",
        "B 级剧情侵入：发一张异常消息卡或在日志末尾追加一行；让它直接呼应当前证据。",
        "C 级事件：每章至多一次；事前确认参与者接受闪帧/突发音效。减弱档改为静态异常，关闭档只叙述结果。",
        "C 级后停止计时和追踪，保留 5-10 秒安静，再用 NPC 消息恢复人与人的联系。",
    ])
    add_callout(doc, "安全工具", "桌游局开场约定“暂停”手势或词。任何参与者提出暂停时，主持人立即停止音效和惊吓，直接转为减弱或关闭档，不追问原因。", fill="F1DCD7", accent=RED)

    add_section_title(doc, "08", "主持脚本：序章", page_break=True)
    add_h2(doc, "开场朗读")
    add_callout(doc, "03:17", "葬礼结束不到一天，陈默的收件箱亮了一下。发件人是林薇。系统没有把它归进诈骗，也没有标记账号异常；它只是像一封普通邮件那样，安静地等着被打开。", fill=BLACK, accent=CRT_GREEN, text_color=WHITE)
    add_h2(doc, "材料发放顺序")
    add_numbered(doc, [
        "任务邮件与自动回执。",
        "附件错误页：只显示 MIRROR-17。",
        "正确搜索后发五条结果，其中镜像状态快照给出地址。",
        "玩家 scan 后发镜像节点目录。",
        "get 林薇日志后触发时间戳异常；随后允许安静阅读全文。",
        "两份证据齐备后发结案报告。",
    ])
    add_h2(doc, "提示卡")
    add_table(doc, ["级别", "台词"], [
        ("1", "附件报错不是乱码，它像一个旧服务编号。"),
        ("2", "邮件里提到的旧邮局，和 MIRROR-17 也许属于同一件事。"),
        ("3", "把“零时邮局”和“MIRROR-17”一起放进搜索框。"),
    ], [1100, 8640], first_col_bold=True)
    add_body(doc, "结案后朗读：陈默保存了日志。屏幕右下角的时间恢复正常，笔记本里却多出一个他不记得输入过的词：REN-0。", bold_prefix="结案后朗读：")

    add_section_title(doc, "09", "主持脚本：01 凌晨两点的门禁", page_break=True)
    add_h2(doc, "开场朗读")
    add_callout(doc, "02:00", "唐遥把四张门禁截图排在一起。四个日期，四个完全相同的时间：02:00:00。沈栀已经失联，学籍页面显示“自愿退学”，但一张被删除的卡每天都在封闭实验室门口准时活过来。", fill=BLACK, accent=CRT_GREEN, text_color=WHITE)
    add_h2(doc, "关键控场点")
    add_table(
        doc,
        ["节点", "必发材料", "事件"],
        [
            ("公开门户", "学籍异动公示。", "让玩家先接受“自愿退学”的官方叙述。"),
            ("N17 缓存", "access 日志、维护记录。", "取得日志后追加 CHENMO / 当前房间 / 02:00。"),
            ("沈栀云盘", "原始论文、0712 录音。", "录音后停顿，再发已注销账号的异常消息。"),
            ("版本网关", "NIGHT-0200 与版本日志。", "首次完整反追踪教学。"),
            ("审计归档", "跨系统清除授权。", "把 REN-0 从神秘名字升级为权限签名。"),
        ],
        [1900, 3200, 4640],
        first_col_bold=True,
    )
    add_h2(doc, "密码与搜索裁定")
    add_bullets(doc, [
        "设备搜索接受：N17 门禁 / GATE-N17 / N17 维护。",
        "云盘：S20417 / df22023。玩家若准确说明 DF2 相机 + 2023 入学年，可视为推理成立。",
        "负责人搜索接受：杜承远 兰圃实验室 / 兰圃 数据实验室 杜承远。",
    ])
    add_body(doc, "结案后朗读：真正走进实验室的不是沈栀，而是她的身份。她本人仍然活着，正在把原始材料交给警方；可批准清除她的人，仍在更高一层系统里。", bold_prefix="结案后朗读：")

    add_section_title(doc, "10", "主持脚本：02 死者在线", page_break=True)
    add_h2(doc, "开场朗读")
    add_callout(doc, "23:47", "苏晚没有先发文字。她发来一段 47 秒的录屏：渡鸦的直播间亮起红点，镜头里是哥哥的房间。问题是，她亲手整理过那间房，也亲手参加过哥哥三个月前的葬礼。", fill=BLACK, accent=RED, text_color=WHITE)
    add_h2(doc, "材料与反转顺序")
    add_numbered(doc, [
        "先给动态样本，让玩家自己做文风与时间分界。",
        "NAS 合同推翻“盗号”，记账补上苏野为妹妹凑学费的动机。",
        "边缘服务器用 IP、编码器与镜面位置推翻“从家中直播”。",
        "工单和排班推翻“录像或 AI”，确认真人扮演与棚 B。",
        "身份池把商业犯罪连接到被系统抹除的活人，并再次出现 REN-0。",
    ])
    add_h2(doc, "密码与搜索裁定")
    add_table(doc, ["门", "有效输入"], [
        ("旧 NAS", "suye / mantou0713；或 raven / MANTOU0713；或 苏野 / 馒头0713。"),
        ("公司搜索", "长明 / 长明网络科技 / changming / 数字纪念馆。"),
        ("工单系统", "CM-041 / 202303changming。"),
    ], [1800, 7940], first_col_bold=True)
    add_h2(doc, "第 39 格")
    add_body(doc, "完整档：在玩家宣布前进到 39/47 时，出示异常帧 0.4 秒并播放单次 sting；收走图片或切黑，保持五秒静默。随后只发苏晚消息：“你还在吗？”", bold_prefix="完整档：")
    add_body(doc, "减弱档：直接出示静态异常并说“镜中人已经转头”；无闪帧、无 sting。关闭档：叙述镜面证据，不制造突发刺激。", bold_prefix="减弱档：")
    add_section_title(doc, "11", "结案、后续与创作检查", page_break=True)
    add_h2(doc, "阶段 1 收束")
    add_table(
        doc,
        ["案件", "表层结案", "暗线碎片"],
        [
            ("序章", "延迟邮件而非死者复活。", "REN-0 是被多人用过的钥匙。"),
            ("01", "维护盒重放身份，杜承远栽赃并抹除沈栀。", "REN-0 批准跨学籍、门禁、财务的清除。"),
            ("02", "长明用脚本与真人扮演交易死者账号。", "身份池买方是查无此人的活人，经办标记 REN-0。"),
        ],
        [1500, 4700, 3540],
        first_col_bold=True,
    )
    add_h2(doc, "主持复盘问题")
    add_checklist(doc, [
        "玩家是否亲手说出了关键搜索词与密码构成？",
        "每个重大结论是否至少有两份不同来源的材料支持？",
        "恐怖事件是否紧贴当前真相，而不是随机打断？",
        "惊吓后是否留出安静和 NPC 人性回应？",
        "结案是否明确违法有代价、材料将进入警方调查？",
        "是否把 REN-0 保持为跨系统机制线索，而不是提前具体化成终极反派？",
    ])
    add_h2(doc, "阶段 2 方向")
    add_bullets(doc, [
        "03《午夜的水表》：健康档案被远程改写，水电数据证明死者家中有第二个人；恐怖浓度最高。",
        "04《被优化的人》：失业、拒贷、社交蒸发组成信用降级链，暗线浮出水面。",
        "05《名单上的名字》：林薇、历关受害者与陈默都在名单上，锁定默许机制的人。",
        "06《收网》：整合证据、自首、配合警方收网，完成责任与正义主题。",
    ], compact=True)
    add_section_title(doc, "12", "主持速查页", page_break=True)
    add_table(
        doc,
        ["项目", "序章", "01 门禁", "02 死者在线"],
        [
            ("知识门", "零时邮局 + MIRROR-17", "N17 门禁；DF2 + 2023；杜承远 + 兰圃", "馒头 + 0713；长明；CM-041 规则"),
            ("核心证据", "林薇日志 / 投递日志", "门禁日志 / 维护记录 / 录音 / 清除授权", "合同 / 记账 / 推流日志 / 排班 / 身份池"),
            ("追踪", "无", "版本网关 120 秒；审计归档 90 秒", "边缘 120 秒；长明 150 秒；棚区 90 秒"),
            ("恐怖", "未来时间戳", "日志写入陈默；已注销账号消息", "新动态；私信；第 39 格；trace 陈默"),
            ("REN-0", "钥匙", "跨系统批准签名", "身份池经办标记"),
        ],
        [1500, 2500, 2870, 2870],
        first_col_bold=True,
        font_size=8.5,
    )
    add_code_block(doc, [
        "HOST RULE 01 : REVEAL ORIGINALS, NOT CONCLUSIONS",
        "HOST RULE 02 : ACCEPT REASONABLE ALIASES",
        "HOST RULE 03 : ONE C-LEVEL EVENT MAXIMUM PER CASE",
        "HOST RULE 04 : PAUSE MEANS PAUSE",
        "HOST RULE 05 : END WITH EVIDENCE, RESPONSIBILITY, POLICE",
    ], label="FACILITATOR CHECKSUM")

    doc.save(WORLD_BOOK)


if __name__ == "__main__":
    build_player_guide()
    build_world_book()
    print(PLAYER_GUIDE)
    print(WORLD_BOOK)
